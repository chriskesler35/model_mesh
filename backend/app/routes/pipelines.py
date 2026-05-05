"""Multi-agent pipeline orchestration (Option A).

A Pipeline runs a task through N specialist phases (Analyst → Architect →
Coder → Reviewer → ...). Each phase is one LLM call with its own role,
system prompt, and structured artifact output. Phases gate on user approval
between runs (unless auto_approve=True).

Architecture:
  - One workbench_session owns the pipeline (via session.pipeline_id)
  - Each phase is a PhaseRun row (persisted for audit + replay)
  - Prior phase artifacts are injected as context for downstream phases
  - SSE events stream to the frontend via the pipeline's own event queue
  - Approval endpoints advance the pipeline; rejection re-runs the current phase
"""

import asyncio
import json
import logging
import re
import uuid
from datetime import datetime
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Optional, Dict, Any, List, AsyncGenerator

from fastapi import APIRouter, HTTPException, Request, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy import select, desc
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import BaseModel

from app.middleware.auth import verify_api_key
from app.database import get_db, AsyncSessionLocal

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/v1/workbench/pipelines", tags=["pipelines"])


# ─── In-memory event queues for SSE (state lives in DB) ───────────────────────
_queues:     Dict[str, asyncio.Queue] = {}
_event_logs: Dict[str, List[dict]]    = {}


# ─── Schemas ──────────────────────────────────────────────────────────────────
class PipelineCreate(BaseModel):
    session_id: str
    method_id: str                              # e.g. bmad | gsd | superpowers | specaudit | mvp-loop | custom_method_id
    stack_override: Optional[List[str]] = None  # explicit launcher-selected stack order
    task: str
    auto_approve: bool = False
    interaction_mode: str = "autonomous"               # 'interactive' | 'autonomous'
    delegate_qa_to_agent: bool = False                  # autonomous surrogate for user Q&A
    model_overrides: Optional[Dict[str, str]] = None   # {phase_name: model_id}
    approvers: Optional[List[str]] = None              # list of user IDs
    approval_policy: Optional[str] = "any"             # 'any', 'majority', 'all'


class PipelineApprove(BaseModel):
    feedback: Optional[str] = None
    user_id: Optional[str] = None              # approver user ID (optional, falls back to request user)


class PipelineReject(BaseModel):
    feedback: str                                # required — why we're rejecting


class PipelineSkip(BaseModel):
    reason: Optional[str] = None


class PipelineMessage(BaseModel):
    message: str


class PipelinePhaseModelUpdate(BaseModel):
    phase_index: int
    model_id: str


class PipelinePhasePreviewRequest(BaseModel):
    method_id: str
    stack_override: Optional[List[str]] = None


def _apply_interaction_mode_to_phases(
    phases: List[Dict[str, Any]],
    interaction_mode: str,
    delegate_qa_to_agent: bool,
) -> List[Dict[str, Any]]:
    """Inject interaction-mode instructions into method phases.

    - interactive: force requirements discovery output to include explicit
      unanswered questions for the human user.
    - autonomous + delegate_qa_to_agent: produce a visible surrogate Q&A block
      with inferred user answers and confidence levels.
    """
    mode = (interaction_mode or "autonomous").strip().lower()
    delegated = bool(delegate_qa_to_agent)

    if mode not in ("interactive", "autonomous"):
        raise HTTPException(status_code=400, detail="interaction_mode must be 'interactive' or 'autonomous'")

    updated: List[Dict[str, Any]] = []
    # Apply discovery guidance to kickoff/root phases across all methods,
    # including custom methods, so behavior matches in-IDE kickoff patterns.
    root_phase_names = {
        ph.get("name")
        for ph in phases
        if not (ph.get("depends_on") or [])
    }

    for phase in phases:
        p = dict(phase)
        p["interaction_mode"] = mode
        p["delegate_qa_to_agent"] = delegated

        if p.get("name") in root_phase_names:
            base_prompt = p.get("system_prompt", "")
            artifact_type = (p.get("artifact_type") or "md").lower()
            if mode == "interactive":
                if artifact_type == "json":
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: INTERACTIVE USER DISCOVERY\n"
                        "You must surface a clear Q&A set for the human reviewer before implementation proceeds.\n"
                        "Do not invent missing user decisions. Mark unknowns explicitly.\n\n"
                        "Add these fields to your JSON output:\n"
                        "- clarifying_questions: [{id, question, why_it_matters}]\n"
                        "- open_questions_for_user: [id]\n"
                        "- assumed_answers: []\n"
                    )
                elif artifact_type == "code":
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: INTERACTIVE USER DISCOVERY\n"
                        "Before any FILE: or CMD: blocks, include a short DISCOVERY section in markdown with:\n"
                        "- Clarifying questions (numbered)\n"
                        "- Open questions for user\n"
                        "- Explicitly empty assumed answers list\n"
                        "If critical requirements are unknown, keep implementation minimal/safe and call out the blockers clearly.\n"
                    )
                else:  # md/other
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: INTERACTIVE USER DISCOVERY\n"
                        "Your output must include these sections:\n"
                        "## Clarifying Questions\n"
                        "## Open Questions For User\n"
                        "## Assumed Answers (leave empty in interactive mode)\n"
                    )
            elif delegated:
                if artifact_type == "json":
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: AUTONOMOUS SURROGATE Q&A\n"
                        "Generate visible Q&A as if an end user answered, so the conversation remains auditable.\n"
                        "Use reasonable assumptions and label confidence clearly.\n\n"
                        "Add these fields to your JSON output:\n"
                        "- clarifying_questions: [{id, question, why_it_matters}]\n"
                        "- assumed_answers: [{question_id, answer, confidence: high|medium|low, rationale}]\n"
                        "- open_questions_for_user: [id]  # only if uncertainty is too high\n"
                    )
                elif artifact_type == "code":
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: AUTONOMOUS SURROGATE Q&A\n"
                        "Before any FILE: or CMD: blocks, include a DISCOVERY section with:\n"
                        "- Clarifying questions\n"
                        "- Assumed answers with confidence (high|medium|low) and rationale\n"
                        "- Open questions for user only when uncertainty is too high\n"
                    )
                else:  # md/other
                    p["system_prompt"] = (
                        f"{base_prompt}\n\n"
                        "# Interaction mode: AUTONOMOUS SURROGATE Q&A\n"
                        "Your output must include these sections:\n"
                        "## Clarifying Questions\n"
                        "## Assumed Answers (with confidence + rationale)\n"
                        "## Open Questions For User (only high-uncertainty blockers)\n"
                    )

        updated.append(p)

    return updated


def _resolve_pipeline_method_selection(
    method_id: str,
    stack_override: Optional[List[str]] = None,
) -> tuple[str, List[str], str, List[str]]:
    """Resolve effective method + stacked prompt behavior exactly like create_pipeline."""
    from app.routes.methods import _load_state as _load_method_state, BUILT_IN_METHODS, _build_stack_prompt

    effective_method_id = method_id
    stacked_prompt = ""
    requested_stack = [m for m in (stack_override or []) if m != "standard"]
    invalid_stack = [m for m in requested_stack if m not in BUILT_IN_METHODS]
    if invalid_stack:
        raise HTTPException(status_code=400, detail=f"Unknown methods in stack_override: {invalid_stack}")

    if method_id in ("stack", "active"):
        method_state = _load_method_state()
        stack = requested_stack or method_state.get("active_stack", [])
        primary = stack[0] if stack else method_state.get("active_method", "standard")
        effective_method_id = primary
        non_primary = [m for m in stack if m != primary]
        if non_primary:
            stacked_prompt = _build_stack_prompt(non_primary)
    elif method_id in ("bmad", "gsd", "superpowers", "specaudit", "mvp-loop"):
        method_state = _load_method_state()
        stack = requested_stack or method_state.get("active_stack", [])
        if len(stack) > 1 and method_id in stack:
            non_primary = [m for m in stack if m != method_id]
            if non_primary:
                stacked_prompt = _build_stack_prompt(non_primary)

    layered_methods = []
    if stacked_prompt:
        if method_id in ("stack", "active"):
            layered_methods = [m for m in requested_stack if m != effective_method_id]
        else:
            layered_methods = [m for m in requested_stack if m != method_id]
    return effective_method_id, requested_stack, stacked_prompt, layered_methods


async def _build_phase_preview_rows(method_id: str) -> List[Dict[str, Any]]:
    """Build phase preview rows with resolved agent/persona/model chain."""
    from app.services.phase_templates import get_phases_for_method, list_supported_methods
    from app.models.agent import Agent
    from app.models.persona import Persona
    from app.routes.workbench import _resolve_model

    try:
        phases = get_phases_for_method(method_id)
    except KeyError:
        raise HTTPException(
            status_code=404,
            detail=f"Method '{method_id}' not supported. Available: {list_supported_methods()}"
        )

    async with AsyncSessionLocal() as db:
        agent_rows = (await db.execute(
            select(Agent).where(Agent.is_active == True).where(Agent.method_phase.isnot(None))
        )).scalars().all()
        agent_by_phase = {}
        for a in agent_rows:
            key = (a.method_phase or "").lower()
            if key and key not in agent_by_phase:
                agent_by_phase[key] = a

        persona_rows = (await db.execute(select(Persona))).scalars().all()
        persona_by_id = {str(p.id): p for p in persona_rows}
        persona_by_name = {p.name.lower(): p for p in persona_rows}

    async def chain_info(phase_name: str) -> dict:
        result = {
            "has_agent": False, "agent_name": None,
            "has_persona": False, "persona_name": None,
            "resolved_model": None, "resolved_via": None,
        }
        agent = agent_by_phase.get(phase_name.lower())
        if agent:
            result["has_agent"] = True
            result["agent_name"] = agent.name
            if agent.persona_id:
                persona = persona_by_id.get(str(agent.persona_id))
                if persona:
                    result["has_persona"] = True
                    result["persona_name"] = persona.name
                    if persona.primary_model_id:
                        m, _ = await _resolve_model(str(persona.primary_model_id))
                        if m:
                            result["resolved_model"] = m.model_id
                            result["resolved_via"] = "agent→persona"
                            return result
            if agent.model_id:
                m, _ = await _resolve_model(str(agent.model_id))
                if m:
                    result["resolved_model"] = m.model_id
                    result["resolved_via"] = "agent→model"
                    return result

        persona = persona_by_name.get(phase_name.lower())
        if persona:
            result["has_persona"] = True
            result["persona_name"] = persona.name
            if persona.primary_model_id:
                m, _ = await _resolve_model(str(persona.primary_model_id))
                if m:
                    result["resolved_model"] = m.model_id
                    result["resolved_via"] = "persona name-match"
                    return result
        return result

    phase_rows = []
    for p in phases:
        info = await chain_info(p["name"])
        phase_rows.append(
            {
                "name": p["name"],
                "role": p["role"],
                "default_model": p["default_model"],
                "artifact_type": p["artifact_type"],
                "depends_on": p.get("depends_on", []),
                **info,
                "persona_model": None,
            }
        )
    return phase_rows


# ─── Event helpers ────────────────────────────────────────────────────────────
def _push(pipeline_id: str, type: str, **payload):
    evt = {"type": type, "payload": payload, "ts": datetime.utcnow().isoformat()}
    log = _event_logs.setdefault(pipeline_id, [])
    log.append(evt)
    if len(log) > 500:
        del log[0:len(log) - 500]
    q = _queues.get(pipeline_id)
    if q:
        try:
            q.put_nowait(evt)
        except asyncio.QueueFull:
            pass


async def _db_update_pipeline(pipeline_id: str, **kwargs):
    from app.models.pipeline import Pipeline
    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))
            p = result.scalar_one_or_none()
            if p:
                for k, v in kwargs.items():
                    setattr(p, k, v)
                await db.commit()
    except Exception as e:
        logger.warning(f"Pipeline DB update failed for {pipeline_id}: {e}")


async def _db_update_phase(phase_run_id: str, **kwargs):
    from app.models.pipeline import PhaseRun
    try:
        async with AsyncSessionLocal() as db:
            result = await db.execute(select(PhaseRun).where(PhaseRun.id == phase_run_id))
            pr = result.scalar_one_or_none()
            if pr:
                for k, v in kwargs.items():
                    setattr(pr, k, v)
                await db.commit()
    except Exception as e:
        logger.warning(f"PhaseRun DB update failed for {phase_run_id}: {e}")


async def _notify_approvers(pipeline_id: str, phase_index: int, phase_name: str):
    """Send notifications + WebSocket alerts to all pipeline approvers when a phase needs approval."""
    from app.models.pipeline import Pipeline
    from app.models.notification import Notification
    from app.services.ws_manager import manager

    try:
        async with AsyncSessionLocal() as db:
            p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
            if not p:
                return

            approvers = p.approvers or []
            if not approvers:
                # Fallback to pipeline creator
                approvers = [p.created_by] if p.created_by else ["owner"]

            for user_id in approvers:
                notif = Notification(
                    user_id=user_id,
                    type="pipeline_approval",
                    title=f"Approval needed: {phase_name}",
                    message=f"Phase '{phase_name}' in pipeline is awaiting your approval.",
                    conversation_id=pipeline_id,
                )
                db.add(notif)

                # Real-time WebSocket push
                try:
                    await manager.send_to_user(user_id, {
                        "type": "pipeline_approval_needed",
                        "payload": {
                            "pipeline_id": pipeline_id,
                            "phase_index": phase_index,
                            "phase_name": phase_name,
                            "approval_policy": p.approval_policy or "any",
                        },
                    })
                except Exception as ws_err:
                    logger.debug(f"WebSocket notify failed for {user_id}: {ws_err}")

            await db.commit()
    except Exception as e:
        logger.warning(f"Failed to notify approvers for pipeline {pipeline_id}: {e}")


def _check_approval_threshold(approvals: list, approvers: list, policy: str) -> str:
    """Check if approval threshold is met based on policy.

    Returns: 'approved', 'rejected', or 'pending'.
    Any single rejection immediately rejects regardless of policy.
    """
    if not approvals:
        return "pending"

    approve_count = sum(1 for a in approvals if a.get("action") == "approve")
    reject_count = sum(1 for a in approvals if a.get("action") == "reject")

    # Any rejection immediately rejects
    if reject_count > 0:
        return "rejected"

    total_approvers = max(len(approvers), 1)

    if policy == "any":
        return "approved" if approve_count >= 1 else "pending"
    elif policy == "majority":
        needed = (total_approvers // 2) + 1
        return "approved" if approve_count >= needed else "pending"
    elif policy == "all":
        return "approved" if approve_count >= total_approvers else "pending"

    # Default to 'any'
    return "approved" if approve_count >= 1 else "pending"


# ─── Artifact extractors ──────────────────────────────────────────────────────
def _extract_json_artifact(text: str) -> Optional[dict]:
    """Pull the first ```json ... ``` block out of an LLM response."""
    m = re.search(r'```json\s*\n(.*?)```', text, re.DOTALL)
    if not m:
        # fallback — any fenced block that parses as JSON
        m = re.search(r'```[a-z]*\s*\n(\{.*?\}|\[.*?\])\s*```', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception as e:
            logger.debug(f"JSON parse failed: {e}")
            # Try to find a bare JSON object
            try:
                m2 = re.search(r'(\{[\s\S]*\})', text)
                if m2:
                    return json.loads(m2.group(1))
            except Exception:
                pass
    return None


def _extract_code_files(text: str) -> List[dict]:
    """Reuse the workbench FILE: parser to extract code files."""
    from app.routes.workbench import _parse_files
    return _parse_files(text)


def _json_value_to_markdown(value: Any, level: int = 2, name: Optional[str] = None) -> str:
    """Render JSON data into readable markdown for human review.

    The structured JSON remains the source of truth for downstream phases.
    This markdown companion exists purely for approver readability.
    """
    heading = "#" * max(2, min(level, 6))

    if isinstance(value, dict):
        parts: List[str] = []
        if name:
            parts.append(f"{heading} {name}")
        if not value:
            parts.append("(empty object)")
            return "\n\n".join(parts)

        scalar_lines: List[str] = []
        nested_parts: List[str] = []
        for key, item in value.items():
            label = str(key).replace("_", " ").strip().title()
            if isinstance(item, (dict, list)):
                nested_parts.append(_json_value_to_markdown(item, level + 1, label))
            else:
                scalar_lines.append(f"- **{label}:** {item}")

        if scalar_lines:
            parts.append("\n".join(scalar_lines))
        if nested_parts:
            parts.extend(part for part in nested_parts if part)
        return "\n\n".join(part for part in parts if part)

    if isinstance(value, list):
        parts = []
        if name:
            parts.append(f"{heading} {name}")
        if not value:
            parts.append("(empty list)")
            return "\n\n".join(parts)

        simple_values = all(not isinstance(item, (dict, list)) for item in value)
        if simple_values:
            parts.append("\n".join(f"- {item}" for item in value))
            return "\n\n".join(parts)

        rendered_items: List[str] = []
        for idx, item in enumerate(value, start=1):
            item_name = f"Item {idx}"
            rendered_items.append(_json_value_to_markdown(item, level + 1, item_name))
        parts.extend(rendered_items)
        return "\n\n".join(part for part in parts if part)

    if name:
        return f"{heading} {name}\n\n{value}"
    return str(value)


def _json_artifact_review_markdown(data: Any) -> str:
    if data is None:
        return "## Review Preview\n\nNo structured JSON content was parsed from this phase."
    rendered = _json_value_to_markdown(data, level=2)
    return rendered or "## Review Preview\n\nNo structured JSON content was parsed from this phase."


def _build_artifact(phase_def: dict, raw_response: str) -> dict:
    """Convert raw LLM text into the structured artifact for this phase type."""
    atype = phase_def.get("artifact_type", "md")
    if atype == "json":
        data = _extract_json_artifact(raw_response)
        return {
            "type": "json",
            "data": data,
            "review_markdown": _json_artifact_review_markdown(data),
            "raw": raw_response,
        }
    elif atype == "code":
        files = _extract_code_files(raw_response)
        return {"type": "code", "files": files, "raw": raw_response}
    else:  # md
        return {"type": "md", "content": raw_response.strip(), "raw": raw_response}


async def _persist_retrospective_memory_if_applicable(pipeline_id: str) -> None:
    """Persist the MemoryCurator output for retrospective pipelines into user memory."""
    from app.models.pipeline import Pipeline, PhaseRun
    from app.services.memory_context import MemoryContext

    async with AsyncSessionLocal() as db:
        pipeline = (await db.execute(
            select(Pipeline).where(Pipeline.id == pipeline_id)
        )).scalar_one_or_none()
        if not pipeline or pipeline.method_id != "retrospective":
            return

        final_run = (await db.execute(
            select(PhaseRun)
            .where(PhaseRun.pipeline_id == pipeline_id)
            .where(PhaseRun.phase_name == "MemoryCurator")
            .order_by(desc(PhaseRun.created_at))
        )).scalars().first()
        if not final_run:
            return

        artifact = final_run.output_artifact or {}
        if artifact.get("type") != "json":
            return

        data = artifact.get("data") or {}
        memory_markdown = (data.get("memory_markdown") or "").strip()
        if not memory_markdown:
            return

        memory = MemoryContext(db)
        files = await memory.get_memory_files()
        existing = (files.get("RETROSPECTIVES.md") or "").strip()
        content = f"{existing}\n\n---\n\n{memory_markdown}".strip() if existing else memory_markdown
        await memory.update_memory_file("RETROSPECTIVES.md", content)


def _format_prior_artifact_for_context(phase_run_dict: dict) -> str:
    """Render a prior phase's artifact as context for downstream phases."""
    phase_name = phase_run_dict.get("phase_name", "?")
    role = phase_run_dict.get("agent_role", "?")
    artifact = phase_run_dict.get("output_artifact") or {}
    atype = artifact.get("type")

    header = f"## Prior Phase: {phase_name} ({role})"
    if atype == "json":
        data = artifact.get("data")
        if data:
            return f"{header}\n```json\n{json.dumps(data, indent=2)}\n```"
        return f"{header}\n(no structured output — raw text below)\n{artifact.get('raw', '')[:2000]}"
    elif atype == "code":
        files = artifact.get("files", []) or []
        if not files:
            return f"{header}\n(no files generated)"
        file_list = "\n".join(f"- {f['path']}" for f in files)
        previews = "\n\n".join(
            f"### {f['path']}\n```\n{f['content'][:1500]}\n```"
            for f in files[:5]
        )
        return f"{header}\nFiles written:\n{file_list}\n\n{previews}"
    else:  # md
        return f"{header}\n{artifact.get('content', '')[:3000]}"


# ─── Parallel context merging (E2.4) ──────────────────────────────────────────
# Default context limit in characters (~100k chars ≈ ~25k tokens)
DEFAULT_CONTEXT_CHAR_LIMIT = 100_000


def _merge_parent_contexts(
    pipeline_phases: List[dict],
    phase_name: str,
    prior_run_dicts: List[dict],
) -> str:
    """Merge artifacts from all parent phases into a single context block.

    Detects artifact types and applies type-appropriate merge strategies:
    - JSON artifacts → merged into a single JSON object keyed by phase name
    - Code artifacts → concatenated with file-boundary markers
    - Text/md artifacts → concatenated with section headers
    - Mixed types → each rendered with its own strategy, then concatenated
    """
    if not prior_run_dicts:
        return ""

    # Classify artifacts by type
    json_parts: Dict[str, Any] = {}
    code_parts: List[str] = []
    text_parts: List[str] = []

    for prd in prior_run_dicts:
        pname = prd.get("phase_name", "unknown")
        artifact = prd.get("output_artifact") or {}
        atype = artifact.get("type")

        if atype == "json":
            data = artifact.get("data")
            if data is not None:
                json_parts[pname] = data
            else:
                # Fallback: treat raw text as text artifact
                raw = artifact.get("raw", "")
                if raw:
                    text_parts.append(f"## Output from {pname}\n{raw[:3000]}")
        elif atype == "code":
            files = artifact.get("files", []) or []
            if files:
                section = [f"## Output from {pname}"]
                for f in files:
                    path = f.get("path", "unknown")
                    content = f.get("content", "")
                    section.append(
                        f"### FILE: {path}\n```\n{content}\n```"
                    )
                code_parts.append("\n\n".join(section))
            else:
                text_parts.append(f"## Output from {pname}\n(no files generated)")
        else:
            # md / text / unknown
            content = artifact.get("content") or artifact.get("raw", "")
            text_parts.append(f"## Output from {pname}\n{content[:3000]}")

    # Assemble merged context
    sections: List[str] = []

    if json_parts:
        merged_json = json.dumps(json_parts, indent=2)
        sections.append(
            f"## Merged JSON from parent phases\n```json\n{merged_json}\n```"
        )

    if code_parts:
        sections.extend(code_parts)

    if text_parts:
        sections.extend(text_parts)

    return "\n\n---\n\n".join(sections)


def _validate_context_size(
    context: str,
    limit: int = DEFAULT_CONTEXT_CHAR_LIMIT,
) -> str:
    """Validate and truncate merged context if it exceeds the character limit.

    Returns the context (possibly truncated) with a warning header if trimmed.
    """
    if len(context) <= limit:
        return context

    truncated = context[:limit]
    # Try to cut at the last section boundary to keep output coherent
    last_sep = truncated.rfind("\n\n---\n\n")
    if last_sep > limit // 2:
        truncated = truncated[:last_sep]

    warning = (
        f"⚠️ Parent context was truncated from {len(context):,} to "
        f"{len(truncated):,} characters (limit: {limit:,}).\n\n"
    )
    return warning + truncated


# ─── Phase runner (one phase = one LLM call) ──────────────────────────────────
async def _run_phase(pipeline_id: str, phase_index: int, retry_count: int = 0, max_retries: int = 0):
    """Execute a single phase. Persists PhaseRun + artifact. Streams events."""
    from app.models.pipeline import Pipeline, PhaseRun
    from app.services.model_client import ModelClient
    from app.services.identity_context import build_identity_context
    from app.routes.workbench import _build_runtime_model_chain, _humanize_model_error, _should_failover_error, _resolve_model

    # Load pipeline
    async with AsyncSessionLocal() as db:
        p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
        if not p:
            logger.error(f"_run_phase: pipeline {pipeline_id} not found")
            return
        if p.status in ("paused", "cancelled"):
            logger.info("_run_phase: pipeline %s is %s; skipping phase launch", pipeline_id, p.status)
            return
        phases = p.phases or []
        initial_task = p.initial_task
        auto_approve = p.auto_approve

        # Load phase runs from dependency phases only (not all prior phases).
        # This ensures each phase gets context specifically from the phases
        # it declared in depends_on, enabling correct parallel execution.
        all_prior_runs = (await db.execute(
            select(PhaseRun)
            .where(PhaseRun.pipeline_id == pipeline_id)
            .where(PhaseRun.status.in_(("approved", "skipped", "completed")))
            .order_by(PhaseRun.phase_index)
        )).scalars().all()

    if phase_index >= len(phases):
        logger.info(f"Pipeline {pipeline_id}: all phases complete")
        await _db_update_pipeline(pipeline_id, status="completed", completed_at=datetime.utcnow())
        _push(pipeline_id, "pipeline_done", message="All phases complete.", status="completed")
        return

    phase_def = phases[phase_index]
    phase_name = phase_def["name"]

    # ── Handle branch phases — no LLM call, just routing ───────────────────
    if phase_def.get("phase_type") == "branch":
        from app.services.phase_templates import evaluate_branch

        # Get parent output from the most recent prior run
        parent_output = ""
        if prior_run_dicts:
            last = prior_run_dicts[-1]
            artifact = last.get("output_artifact") or {}
            if artifact.get("type") == "json" and artifact.get("data") is not None:
                parent_output = json.dumps(artifact["data"])
            elif artifact.get("raw"):
                parent_output = artifact["raw"]

        target = evaluate_branch(phase_def, parent_output)

        # Create PhaseRun as instantly completed (no LLM, no tokens)
        run_id = str(uuid.uuid4())
        now = datetime.utcnow()
        async with AsyncSessionLocal() as db:
            pr = PhaseRun(
                id=run_id,
                pipeline_id=pipeline_id,
                phase_index=phase_index,
                phase_name=phase_name,
                agent_role="branch",
                model_id=None,
                status="approved",
                started_at=now,
                completed_at=now,
                input_context={
                    "branch_target": target,
                    "parent_output_preview": (parent_output or "")[:200],
                },
                output_artifact={"type": "branch", "target": target},
            )
            db.add(pr)
            await db.commit()

        _push(pipeline_id, "phase_branch",
              phase_index=phase_index,
              phase_name=phase_name,
              target=target)

        logger.info(f"Pipeline {pipeline_id}: branch phase '{phase_name}' routed to '{target}'")

        await _advance_to_next(pipeline_id, phase_index)
        return

    agent_role = phase_def["role"]

    # Filter prior runs to only those from declared dependencies
    depends_on = set(phase_def.get("depends_on") or [])
    if depends_on:
        prior_run_dicts = [
            r.to_dict() for r in all_prior_runs
            if r.phase_name in depends_on
        ]
    else:
        # No explicit deps — for phases with empty depends_on (root phases),
        # no prior context is injected.  For legacy phases without the field,
        # fall back to all prior runs so existing behaviour is preserved.
        if "depends_on" in phase_def:
            prior_run_dicts = []
        else:
            prior_run_dicts = [r.to_dict() for r in all_prior_runs]

    # ── Condition evaluation ─────────────────────────────────────────────
    if phase_def.get("conditions"):
        from app.services.phase_templates import evaluate_phase_conditions, format_condition_reason

        parent_output: Optional[str] = None
        if prior_run_dicts:
            last_parent = prior_run_dicts[-1]
            artifact_data = last_parent.get("output_artifact") or {}
            if artifact_data.get("type") == "json" and artifact_data.get("data") is not None:
                parent_output = json.dumps(artifact_data["data"])
            elif artifact_data.get("raw"):
                parent_output = artifact_data["raw"]

        if not evaluate_phase_conditions(phase_def, parent_output):
            skip_reasons = [format_condition_reason(c) for c in phase_def["conditions"]]
            reason = "; ".join(skip_reasons)

            skip_run_id = str(uuid.uuid4())
            async with AsyncSessionLocal() as db:
                pr = PhaseRun(
                    id=skip_run_id,
                    pipeline_id=pipeline_id,
                    phase_index=phase_index,
                    phase_name=phase_name,
                    agent_role=agent_role,
                    model_id=None,
                    status="skipped",
                    started_at=datetime.utcnow(),
                    completed_at=datetime.utcnow(),
                    input_context={"skip_reason": reason},
                )
                db.add(pr)
                await db.commit()

            logger.info("Pipeline %s: skipping phase %d (%s) — %s", pipeline_id, phase_index, phase_name, reason)
            _push(pipeline_id, "phase_skipped", phase_index=phase_index, phase_name=phase_name, reason=reason)
            await _advance_to_next(pipeline_id)
            return

    # Model resolution chain (highest priority wins):
    #   1. Per-phase user override (phase_def["model"] from model_overrides)
    #   2. Agent bound to this method_phase → its persona → persona.primary_model
    #   3. Agent bound to this method_phase → its own model_id (if no persona)
    #   4. Persona matching phase name (legacy, direct persona match)
    #   5. Template default_model (fallback)
    resolved_model = None
    resolved_system_prompt = None  # persona + agent prompts prepended to phase prompt
    resolved_via = None             # for logging
    explicit_fallback_models: list[str] = []

    if not phase_def.get("model"):
        try:
            from app.models.agent import Agent
            from app.models.persona import Persona
            from sqlalchemy import func as sqlfunc

            async with AsyncSessionLocal() as db:
                # Look up the agent bound to this phase (case-insensitive)
                agent = (await db.execute(
                    select(Agent)
                    .where(sqlfunc.lower(Agent.method_phase) == phase_name.lower())
                    .where(Agent.is_active == True)
                    .order_by(Agent.created_at)
                    .limit(1)
                )).scalar_one_or_none()

                if agent:
                    extra_prompts = []
                    # If the agent has its own system_prompt, prepend it
                    if agent.system_prompt:
                        extra_prompts.append(f"# Agent: {agent.name}\n{agent.system_prompt}")
                    # Resolve model via persona if attached
                    if agent.persona_id:
                        persona = (await db.execute(
                            select(Persona).where(Persona.id == agent.persona_id)
                        )).scalar_one_or_none()
                        if persona:
                            if persona.system_prompt:
                                extra_prompts.append(f"# Persona: {persona.name}\n{persona.system_prompt}")
                            if persona.primary_model_id:
                                primary_resolved, _ = await _resolve_model(str(persona.primary_model_id))
                                if primary_resolved:
                                    resolved_model = primary_resolved.model_id
                                    resolved_via = f"agent '{agent.name}' → persona '{persona.name}'"
                            if persona.fallback_model_id:
                                fallback_resolved, _ = await _resolve_model(str(persona.fallback_model_id))
                                if fallback_resolved:
                                    explicit_fallback_models.append(str(fallback_resolved.id))
                    # Fallback: agent's own model_id if no persona
                    if not resolved_model and agent.model_id:
                        direct_resolved, _ = await _resolve_model(str(agent.model_id))
                        if direct_resolved:
                            resolved_model = direct_resolved.model_id
                            resolved_via = f"agent '{agent.name}' (direct model)"
                    if extra_prompts:
                        resolved_system_prompt = "\n\n".join(extra_prompts)

                # Legacy fallback: persona with a matching name, even if no agent
                if not resolved_model:
                    persona = (await db.execute(
                        select(Persona).where(sqlfunc.lower(Persona.name) == phase_name.lower())
                    )).scalar_one_or_none()
                    if persona and persona.primary_model_id:
                        primary_resolved, _ = await _resolve_model(str(persona.primary_model_id))
                        if primary_resolved:
                            resolved_model = primary_resolved.model_id
                            resolved_via = f"persona '{persona.name}' (name match, no agent)"
                            if persona.system_prompt:
                                resolved_system_prompt = f"# Persona: {persona.name}\n{persona.system_prompt}"
                        if persona.fallback_model_id:
                            fallback_resolved, _ = await _resolve_model(str(persona.fallback_model_id))
                            if fallback_resolved:
                                explicit_fallback_models.append(str(fallback_resolved.id))
        except Exception as e:
            logger.warning(f"Phase '{phase_name}' resolution lookup failed: {e}")

    if resolved_via:
        logger.info(f"Phase '{phase_name}' → {resolved_via} → model '{resolved_model}'")

    model_id = phase_def.get("model") or resolved_model or phase_def.get("default_model") or "claude-sonnet-4-6"
    # Normalize any UUID/model reference to the provider-facing model_id string.
    # This protects phase execution if a caller accidentally passes a DB UUID.
    normalized_model, _ = await _resolve_model(str(model_id))
    if normalized_model:
        model_id = normalized_model.model_id
    persona_system_prompt = resolved_system_prompt  # used below when composing the phase system prompt

    # Resolve max_retries from phase_def if not passed explicitly
    if max_retries == 0:
        max_retries = phase_def.get("max_retries", 0)

    # Create PhaseRun row
    phase_run_id = str(uuid.uuid4())
    started = datetime.utcnow()
    async with AsyncSessionLocal() as db:
        pr = PhaseRun(
            id=phase_run_id,
            pipeline_id=pipeline_id,
            phase_index=phase_index,
            phase_name=phase_name,
            agent_role=agent_role,
            model_id=model_id,
            status="running",
            started_at=started,
            retry_count=retry_count,
            max_retries=max_retries,
            input_context={
                "prior_phases": [r["phase_name"] for r in prior_run_dicts],
                "depends_on": list(depends_on),
            },
        )
        db.add(pr)
        await db.commit()

    await _db_update_pipeline(pipeline_id, status="running", current_phase_index=phase_index)

    # Build messages: identity + persona prompt + phase system prompt + initial task + prior artifacts
    identity_block = build_identity_context(include_method=False)
    phase_prompt = phase_def.get("system_prompt", "")

    system_parts = []
    if identity_block:
        system_parts.append(identity_block)
    if persona_system_prompt:
        # User-defined persona voice/instructions take effect before the phase-specific role
        system_parts.append(f"# Persona\n{persona_system_prompt}")
    system_parts.append(f"# Your role: {agent_role}\n\n{phase_prompt}")
    system_prompt = "\n\n---\n\n".join(system_parts)

    # Load the session's project path so we can inject the current file snapshot.
    # Without this, code phases hallucinate new files instead of editing existing ones.
    session_project_path: Optional[Path] = None
    recent_guidance: list[str] = []
    try:
        from app.models.workbench import WorkbenchSession
        async with AsyncSessionLocal() as db:
            sess = (await db.execute(
                select(WorkbenchSession).where(WorkbenchSession.pipeline_id == pipeline_id)
            )).scalar_one_or_none()
            if sess and sess.project_path:
                session_project_path = Path(sess.project_path)
            if sess and sess.messages:
                recent_guidance = [
                    str(msg.get("content", "")).strip()
                    for msg in (sess.messages or [])
                    if isinstance(msg, dict)
                    and msg.get("kind") == "pipeline_note"
                    and str(msg.get("content", "")).strip()
                ][-5:]
    except Exception as e:
        logger.debug(f"Could not resolve session project_path: {e}")

    project_snapshot = ""
    if session_project_path and session_project_path.exists():
        try:
            from app.routes.workbench import _read_project_snapshot
            project_snapshot = _read_project_snapshot(session_project_path)
        except Exception as e:
            logger.debug(f"Snapshot read failed: {e}")

    # Build user message: task + project snapshot + merged parent artifacts
    user_parts = [f"# Original task\n{initial_task}"]
    if project_snapshot:
        user_parts.append(project_snapshot)
    if recent_guidance:
        guidance = "\n".join(f"- {note}" for note in recent_guidance)
        user_parts.append(
            "# User guidance\n"
            "Keep these latest notes from the user in mind while you work:\n"
            f"{guidance}"
        )

    # Merge parent phase contexts (E2.4)
    merged_context = _merge_parent_contexts(phases, phase_name, prior_run_dicts)
    if merged_context:
        merged_context = _validate_context_size(merged_context)
        user_parts.append(merged_context)

    user_parts.append(f"# Your turn: {phase_name}\nProduce your artifact per the instructions in your system prompt.")
    user_message = "\n\n---\n\n".join(user_parts)

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
    ]

    # Run LLM
    client = ModelClient()
    model_chain, preflight_note = await _build_runtime_model_chain(
        model_id,
        explicit_fallback_refs=explicit_fallback_models,
    )
    if not model_chain:
        msg = (
            f"Could not resolve model '{model_id}' for phase {phase_name}. "
            "The model may be inactive, not live-validated, missing, or its provider has no working credentials, "
            "and no validated fallback model was available. Revalidate it on the Models page or pick a different model for this phase."
        )
        logger.error(msg)
        await _db_update_phase(
            phase_run_id,
            status="failed",
            raw_response=msg,
            completed_at=datetime.utcnow(),
        )
        _push(
            pipeline_id,
            "phase_failed",
            phase_index=phase_index,
            phase_name=phase_name,
            model_id=model_id,
            error=msg,
        )
        await _advance_to_next(pipeline_id)
        return

    model_orm = None
    provider_orm = None
    stream = None
    last_error = None
    active_model_id = model_id

    if preflight_note:
        _push(pipeline_id, "info", phase_index=phase_index, phase_name=phase_name, message=preflight_note)

    for idx, (candidate_model, candidate_provider, reason) in enumerate(model_chain):
        if idx > 0:
            previous_model = model_chain[idx - 1][0]
            previous_error = _humanize_model_error(str(last_error), previous_model.model_id) if last_error else "The previous model was unavailable."
            _push(
                pipeline_id,
                "info",
                phase_index=phase_index,
                phase_name=phase_name,
                message=(
                    f"Model failover: switched from '{previous_model.model_id}' to '{candidate_model.model_id}' "
                    f"using {reason}. Reason: {previous_error}"
                ),
            )
        try:
            stream = await client.call_model(
                model=candidate_model,
                provider=candidate_provider,
                messages=messages,
                stream=True,
                temperature=0.3,
                max_tokens=16000,
            )
            model_orm = candidate_model
            provider_orm = candidate_provider
            active_model_id = candidate_model.model_id
            break
        except Exception as e:
            last_error = e
            logger.error(f"LLM call setup failed in phase {phase_name} of pipeline {pipeline_id} using {candidate_model.model_id}: {e}")
            if idx + 1 >= len(model_chain) or not _should_failover_error(e):
                break

    if not model_orm or not provider_orm or stream is None:
        msg = _humanize_model_error(str(last_error) if last_error else "", model_id)
        await _db_update_phase(
            phase_run_id,
            status="failed",
            raw_response=msg,
            completed_at=datetime.utcnow(),
        )
        _push(
            pipeline_id,
            "phase_failed",
            phase_index=phase_index,
            phase_name=phase_name,
            model_id=model_id,
            error=msg,
        )
        await _advance_to_next(pipeline_id)
        return

    input_context = {
        "prior_phases": [r["phase_name"] for r in prior_run_dicts],
        "depends_on": list(depends_on),
    }
    if active_model_id != model_id:
        input_context["requested_model"] = model_id
        input_context["resolved_model"] = active_model_id
        await _db_update_phase(phase_run_id, model_id=active_model_id, input_context=input_context)

    _push(
        pipeline_id, "phase_started",
        phase_index=phase_index,
        phase_name=phase_name,
        agent_role=agent_role,
        phase_run_id=phase_run_id,
        model_id=active_model_id,
    )

    full_response = ""
    input_tokens = 0
    output_tokens = 0
    llm_success = True
    llm_error = None

    _push(pipeline_id, "phase_thinking", phase_index=phase_index,
          message=f"{agent_role} is working with {model_orm.display_name or model_orm.model_id}…")

    try:
        chunk_count = 0
        async for chunk in stream:
            delta = ""
            try:
                if hasattr(chunk, "choices") and chunk.choices:
                    delta = chunk.choices[0].delta.content or ""
                elif isinstance(chunk, dict):
                    delta = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
            except Exception:
                pass

            try:
                usage = None
                if hasattr(chunk, "usage") and chunk.usage:
                    usage = chunk.usage
                elif isinstance(chunk, dict) and chunk.get("usage"):
                    usage = chunk["usage"]
                if usage:
                    input_tokens = (getattr(usage, "prompt_tokens", 0)
                                    or (usage.get("prompt_tokens", 0) if isinstance(usage, dict) else 0))
                    output_tokens = (getattr(usage, "completion_tokens", 0)
                                     or (usage.get("completion_tokens", 0) if isinstance(usage, dict) else 0))
            except Exception:
                pass

            if delta:
                full_response += delta
                chunk_count += 1
                if chunk_count % 25 == 0:
                    _push(pipeline_id, "phase_progress",
                          phase_index=phase_index,
                          chars=len(full_response))

    except Exception as e:
        err_str = str(e)
        logger.error(f"LLM call failed in phase {phase_name} of pipeline {pipeline_id}: {e}")
        friendly = _humanize_model_error(err_str, model_orm.model_id if model_orm else model_id)
        llm_success = False
        llm_error = friendly

    # Estimate tokens if missing
    if input_tokens == 0:
        input_tokens = client.estimate_tokens(messages, model_orm)
    if output_tokens == 0 and full_response:
        try:
            import tiktoken
            enc = tiktoken.get_encoding("cl100k_base")
            output_tokens = len(enc.encode(full_response))
        except Exception:
            output_tokens = len(full_response) // 4

    # Write request_log for stats
    try:
        from app.models.request_log import RequestLog
        estimated_cost = client.estimate_cost(input_tokens, output_tokens, model_orm)
        async with AsyncSessionLocal() as db:
            log = RequestLog(
                model_id=str(model_orm.id),
                provider_id=str(provider_orm.id),
                input_tokens=input_tokens,
                output_tokens=output_tokens,
                latency_ms=int((datetime.utcnow() - started).total_seconds() * 1000),
                estimated_cost=estimated_cost,
                success=llm_success,
                error_message=llm_error,
            )
            db.add(log)
            await db.commit()
    except Exception as e:
        logger.warning(f"request_log write failed: {e}")

    if not llm_success:
        await _db_update_phase(
            phase_run_id,
            status="failed",
            raw_response=full_response or llm_error,
            completed_at=datetime.utcnow(),
            input_tokens=input_tokens,
            output_tokens=output_tokens,
        )
        _push(
            pipeline_id,
            "phase_failed",
            phase_index=phase_index,
            phase_name=phase_name,
            model_id=model_orm.model_id if model_orm else model_id,
            error=llm_error,
        )
        # Let _advance_to_next decide pipeline-level status — other parallel
        # branches may still be running or runnable.
        await _advance_to_next(pipeline_id)
        return

    # Extract structured artifact
    artifact = _build_artifact(phase_def, full_response)

    # If this is a "code" phase and we have a project_path, write files to disk
    if artifact.get("type") == "code" and artifact.get("files"):
        try:
            from app.models.workbench import WorkbenchSession
            async with AsyncSessionLocal() as db:
                sess = (await db.execute(
                    select(WorkbenchSession).where(WorkbenchSession.pipeline_id == pipeline_id)
                )).scalar_one_or_none()
                if sess and sess.project_path:
                    project_path_for_write = Path(sess.project_path)
                    written = []
                    for f in artifact["files"]:
                        rel = f["path"].lstrip("/\\")
                        abs_path = project_path_for_write / rel
                        try:
                            abs_path.parent.mkdir(parents=True, exist_ok=True)
                            abs_path.write_text(f["content"], encoding="utf-8")
                            written.append(rel)
                        except Exception as e:
                            logger.warning(f"Failed to write {rel}: {e}")
                    artifact["files_written_to_disk"] = written
                    _push(pipeline_id, "files_written", phase_index=phase_index, files=written)
                elif sess and not sess.project_path:
                    # Warn the user — without a project path, code phases produce
                    # artifacts but nothing lands on disk.
                    logger.warning(f"Pipeline {pipeline_id} code phase has no project_path — files NOT saved")
                    _push(pipeline_id, "warning",
                          phase_index=phase_index,
                          message=f"Code phase '{phase_name}' produced {len(artifact['files'])} files but session has no project_path — files won't be saved to disk. Attach a project when creating the session.")
        except Exception as e:
            logger.warning(f"File writing for phase failed: {e}")

    # Execute any CMD: blocks the agent emitted (same 3-tier classifier as workbench)
    cmd_results_for_context: list[str] = []
    try:
        from app.services.command_executor import (
            parse_cmd_blocks, create_command_record, execute_and_record,
            classify_with_project_trust, format_command_for_context,
            get_first_github_token,
        )
        from app.services.command_classifier import CommandTier
        gh_token = get_first_github_token()

        cmds = parse_cmd_blocks(full_response)
        if cmds and session_project_path:
            # Load session bypass + project sandbox mode
            bypass_mode = False
            sandbox_mode = "full"
            proj_path_str = str(session_project_path)
            sess_row = None
            try:
                from app.models.workbench import WorkbenchSession
                async with AsyncSessionLocal() as db:
                    sess_row = (await db.execute(
                        select(WorkbenchSession).where(WorkbenchSession.pipeline_id == pipeline_id)
                    )).scalar_one_or_none()
                    if sess_row:
                        bypass_mode = bool(sess_row.bypass_approvals)
                        proj_id = sess_row.project_id
                        if proj_id:
                            try:
                                import json as _json
                                from pathlib import Path as _P
                                data_dir = _P(__file__).parent.parent.parent.parent / "data"
                                pf = data_dir / "projects.json"
                                if pf.exists():
                                    projects = _json.loads(pf.read_text(encoding="utf-8"))
                                    proj = projects.get(str(proj_id)) or {}
                                    sandbox_mode = proj.get("sandbox_mode") or "full"
                            except Exception:
                                pass
            except Exception:
                pass

            session_id_for_cmds = sess_row.id if sess_row else None
            if session_id_for_cmds:
                _push(pipeline_id, "info", phase_index=phase_index,
                      message=f"Phase emitted {len(cmds)} command(s) to execute")
                for raw_cmd in cmds:
                    tier = classify_with_project_trust(raw_cmd, sandbox_mode, proj_path_str)
                    if tier == CommandTier.BLOCKED:
                        continue
                    # In pipelines we respect the same tiering BUT we can't
                    # pause for approval mid-phase easily. Policy: auto-approve
                    # auto+notice tiers, skip approval tier (record it so user
                    # can approve via workbench UI if they click through).
                    if tier == CommandTier.APPROVAL and not bypass_mode:
                        # Record as pending so it shows up in the workbench log
                        rec_id = await create_command_record(
                            session_id_for_cmds, raw_cmd, tier,
                            pipeline_id=pipeline_id, phase_run_id=phase_run_id,
                            initial_status="pending",
                        )
                        _push(pipeline_id, "command_awaiting_approval",
                              command_id=rec_id, command=raw_cmd, tier=tier.value,
                              phase_index=phase_index)
                        continue
                    # Run inline (auto, notice, or bypass)
                    rec_id = await create_command_record(
                        session_id_for_cmds, raw_cmd, tier,
                        pipeline_id=pipeline_id, phase_run_id=phase_run_id,
                        initial_status="running",
                    )
                    _push(pipeline_id, "command_running",
                          command_id=rec_id, command=raw_cmd, tier=tier.value,
                          phase_index=phase_index)
                    result = await execute_and_record(
                        rec_id, raw_cmd, session_project_path, bypass_used=bypass_mode, github_token=gh_token
                    )
                    _push(pipeline_id, "command_completed", phase_index=phase_index, **result)
                    cmd_results_for_context.append(format_command_for_context(result))

        elif cmds and not session_project_path:
            logger.warning(f"Pipeline {pipeline_id} phase emitted {len(cmds)} CMD: blocks but no project_path — commands skipped")
            _push(pipeline_id, "warning", phase_index=phase_index,
                  message=f"Phase emitted {len(cmds)} commands but session has no project — commands skipped")
    except Exception as e:
        logger.warning(f"Command execution in phase failed: {e}")

    # Append command results to the artifact for downstream phases to see
    if cmd_results_for_context:
        existing_raw = artifact.get("raw", "")
        artifact["raw"] = existing_raw + "\n\n## Commands executed:\n" + "\n\n".join(cmd_results_for_context)

    # ── Auto-retry with feedback: check if this reviewer phase rejected ────────
    # If the phase has on_failure="retry_with_feedback" and the output contains
    # a rejection verdict, automatically re-run the upstream phase with
    # the reviewer's feedback injected as context.
    auto_retry_triggered = False
    if phase_def.get("on_failure") == "retry_with_feedback":
        rejection_verdict = False
        feedback_text = ""
        try:
            artifact_data = artifact.get("data") if artifact.get("type") == "json" else None
            if artifact_data and isinstance(artifact_data, dict):
                verdict = str(artifact_data.get("overall_verdict", artifact_data.get("verdict", ""))).lower().strip()
                if verdict in ("reject", "rejected", "fail", "failed", "needs_changes"):
                    rejection_verdict = True
                    # Extract feedback details from the artifact
                    feedback_parts = []
                    issues = artifact_data.get("issues", [])
                    if issues and isinstance(issues, list):
                        for issue in issues:
                            if isinstance(issue, dict):
                                sev = issue.get("severity", "")
                                desc = issue.get("description", "")
                                suggestion = issue.get("suggestion", "")
                                file_ref = issue.get("file", "")
                                parts = [f"[{sev}]" if sev else ""]
                                if file_ref:
                                    parts.append(f"in {file_ref}")
                                parts.append(desc)
                                if suggestion:
                                    parts.append(f"Suggestion: {suggestion}")
                                feedback_parts.append(" ".join(p for p in parts if p))
                            elif isinstance(issue, str):
                                feedback_parts.append(issue)
                    missing = artifact_data.get("missing_from_spec", [])
                    if missing:
                        feedback_parts.append("Missing from spec: " + ", ".join(str(m) for m in missing))
                    security = artifact_data.get("security_concerns", [])
                    if security:
                        feedback_parts.append("Security concerns: " + ", ".join(str(s) for s in security))
                    gaps = artifact_data.get("gaps", [])
                    if gaps:
                        feedback_parts.append("Gaps: " + ", ".join(str(g) for g in gaps))
                    recommendations = artifact_data.get("recommendations", [])
                    if recommendations:
                        feedback_parts.append("Recommendations: " + ", ".join(str(r) for r in recommendations))

                    feedback_text = "\n".join(feedback_parts) if feedback_parts else f"Reviewer verdict: {verdict}"
        except Exception as e:
            logger.warning(f"Failed to parse reviewer output for auto-retry: {e}")

        if rejection_verdict and feedback_text:
            # Persist this reviewer run as "rejected" (it produced a rejection verdict)
            await _db_update_phase(
                phase_run_id,
                status="rejected",
                raw_response=full_response,
                output_artifact=artifact,
                completed_at=datetime.utcnow(),
                input_tokens=input_tokens,
                output_tokens=output_tokens,
            )
            _push(
                pipeline_id, "phase_completed",
                phase_index=phase_index,
                phase_name=phase_name,
                agent_role=agent_role,
                phase_run_id=phase_run_id,
                artifact=artifact,
                status="rejected",
                input_tokens=input_tokens,
                output_tokens=output_tokens,
            )

            # Determine the upstream phase to retry (the one right before this reviewer)
            upstream_index = phase_def.get("retry_target", phase_index - 1)
            if upstream_index < 0:
                upstream_index = 0

            # Count existing retries for the upstream phase
            retry_max = phase_def.get("max_retries", 2)
            async with AsyncSessionLocal() as db:
                from sqlalchemy import func as sqlfunc
                existing_retries = (await db.execute(
                    select(sqlfunc.count(PhaseRun.id))
                    .where(PhaseRun.pipeline_id == pipeline_id)
                    .where(PhaseRun.phase_index == upstream_index)
                    .where(PhaseRun.retry_count > 0)
                )).scalar() or 0
                # The actual retry count is the number of retry runs already done
                current_retry_count = existing_retries

            auto_retry_triggered = await _retry_upstream_with_feedback(
                pipeline_id=pipeline_id,
                reviewer_phase_index=phase_index,
                upstream_phase_index=upstream_index,
                feedback_text=feedback_text,
                retry_count=current_retry_count,
                max_retries=retry_max,
            )

    # ── Normal flow (no auto-retry triggered) ─────────────────────────────────
    if not auto_retry_triggered:
        # Persist phase run (awaiting_approval or approved-if-auto)
        final_status = "approved" if auto_approve else "awaiting_approval"
        await _db_update_phase(
            phase_run_id,
            status=final_status,
            raw_response=full_response,
            output_artifact=artifact,
            completed_at=datetime.utcnow(),
            input_tokens=input_tokens,
            output_tokens=output_tokens,
        )

        _push(
            pipeline_id, "phase_completed",
            phase_index=phase_index,
            phase_name=phase_name,
            agent_role=agent_role,
            phase_run_id=phase_run_id,
            artifact=artifact,
            status=final_status,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
        )

        if auto_approve:
            # Advance immediately — launch any phases whose deps are now met
            await _advance_to_next(pipeline_id)
        else:
            await _db_update_pipeline(pipeline_id, status="awaiting_approval")
            _push(pipeline_id, "awaiting_approval", phase_index=phase_index,
                  message=f"Phase '{phase_name}' complete — awaiting your approval.")
            # Notify approvers (collaborative approval from E8.5)
            await _notify_approvers(pipeline_id, phase_index, phase_name)


async def _advance_to_next(pipeline_id: str, _current_index: int = -1):
    """Advance pipeline by launching all phases whose dependencies are met.

    Uses the dependency graph (``depends_on`` on each phase) instead of a
    simple sequential index.  All phases whose deps are fully satisfied are
    launched concurrently via ``asyncio.create_task``.

    The ``_current_index`` parameter is accepted for backward-compat with
    existing call-sites but is **not used** for scheduling decisions.
    """
    from app.models.pipeline import Pipeline, PhaseRun
    from app.services.phase_templates import get_ready_phases

    async with AsyncSessionLocal() as db:
        p = (await db.execute(
            select(Pipeline).where(Pipeline.id == pipeline_id)
        )).scalar_one_or_none()
        if not p:
            return
        if p.status in ("paused", "cancelled"):
            return
        phases = p.phases or []
        if not phases:
            await _db_update_pipeline(
                pipeline_id, status="completed", completed_at=datetime.utcnow()
            )
            _push(pipeline_id, "pipeline_done",
                  message="No phases to run.", status="completed")
            return

        # Gather status of every PhaseRun for this pipeline
        phase_runs = (await db.execute(
            select(PhaseRun)
            .where(PhaseRun.pipeline_id == pipeline_id)
            .order_by(PhaseRun.phase_index)
        )).scalars().all()

    # Classify runs by status
    completed_names: set[str] = set()
    running_names: set[str] = set()
    failed_names: set[str] = set()
    for run in phase_runs:
        if run.status in ("completed", "skipped", "approved"):
            completed_names.add(run.phase_name)
        elif run.status in ("running", "awaiting_approval"):
            running_names.add(run.phase_name)
        elif run.status == "failed":
            failed_names.add(run.phase_name)

    # Determine which phases are ready to launch
    ready = get_ready_phases(phases, completed_names)
    # Filter out phases that are already running, completed, or failed.
    # Failed phases must stay failed until the user explicitly retries.
    ready = [ph for ph in ready
             if ph["name"] not in running_names
             and ph["name"] not in completed_names
             and ph["name"] not in failed_names]

    all_phase_names = {ph["name"] for ph in phases}

    if not ready and not running_names:
        # Nothing running and nothing to launch — pipeline is done (or stuck
        # because of failures blocking downstream).
        if completed_names | failed_names >= all_phase_names:
            # Every phase has a terminal status
            final_status = "completed" if not failed_names else "failed"
            # Update current_phase_index to the highest completed index for display
            max_idx = max(
                (i for i, ph in enumerate(phases) if ph["name"] in completed_names),
                default=len(phases) - 1,
            )
            await _db_update_pipeline(
                pipeline_id,
                status=final_status,
                completed_at=datetime.utcnow(),
                current_phase_index=max_idx,
            )
            if final_status == "completed":
                await _persist_retrospective_memory_if_applicable(pipeline_id)
                _push(pipeline_id, "pipeline_done",
                      message="All phases complete.", status="completed")
            else:
                _push(pipeline_id, "pipeline_done",
                      message="Pipeline finished with failures.", status="failed")
        else:
            # Some phases still pending but nothing can run — blocked by
            # failed dependencies.  Mark as failed so users can retry.
            await _db_update_pipeline(
                pipeline_id, status="failed", completed_at=datetime.utcnow()
            )
            _push(pipeline_id, "pipeline_done",
                  message="Pipeline blocked — upstream phase(s) failed.",
                  status="failed")
        return

    if not ready:
        # Phases are still running; nothing new to launch yet — just wait.
        return

    # Launch all ready phases concurrently
    # Update current_phase_index to the lowest ready index (for display)
    first_ready_idx = next(
        i for i, ph in enumerate(phases) if ph["name"] == ready[0]["name"]
    )
    await _db_update_pipeline(
        pipeline_id, current_phase_index=first_ready_idx, status="running"
    )

    for phase in ready:
        phase_index = next(
            i for i, ph in enumerate(phases) if ph["name"] == phase["name"]
        )
        asyncio.create_task(_run_phase(pipeline_id, phase_index))


async def _retry_upstream_with_feedback(
    pipeline_id: str,
    reviewer_phase_index: int,
    upstream_phase_index: int,
    feedback_text: str,
    retry_count: int,
    max_retries: int,
) -> bool:
    """Re-run an upstream phase with reviewer feedback injected.

    Returns True if retry was initiated, False if max_retries exhausted
    (pipeline pauses for manual approval in that case).
    """
    from app.models.pipeline import Pipeline, PhaseRun

    # Check if retries exhausted
    if retry_count >= max_retries:
        logger.info(
            f"Pipeline {pipeline_id}: max retries ({max_retries}) exhausted "
            f"for phase {upstream_phase_index}. Pausing for manual approval."
        )
        await _db_update_pipeline(pipeline_id, status="awaiting_approval")
        _push(
            pipeline_id, "phase_retry_exhausted",
            phase_index=upstream_phase_index,
            phase_name="",  # filled below
            retry_count=retry_count,
            max_retries=max_retries,
            message=f"Phase has been retried {retry_count} time(s) without passing review. Manual approval required.",
        )
        # Update the event with the actual phase name
        async with AsyncSessionLocal() as db:
            p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
            if p and p.phases and upstream_phase_index < len(p.phases):
                phase_name = p.phases[upstream_phase_index].get("name", f"Phase {upstream_phase_index}")
                _push(
                    pipeline_id, "awaiting_approval",
                    phase_index=upstream_phase_index,
                    message=f"Phase '{phase_name}' failed review after {retry_count} retries. Awaiting manual approval.",
                )
        return False

    # Inject feedback into the upstream phase's system prompt
    async with AsyncSessionLocal() as db:
        p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
        if not p:
            return False

        phases = list(p.phases or [])
        if upstream_phase_index >= len(phases):
            return False

        upstream_phase = phases[upstream_phase_index]
        phase_name = upstream_phase.get("name", f"Phase {upstream_phase_index}")

        # Add retry feedback block to the upstream phase's system prompt
        feedback_block = (
            f"\n\n# Retry attempt {retry_count + 1}/{max_retries} — Reviewer feedback from previous attempt:\n"
            f"{feedback_text}\n\n"
            f"Address ALL the reviewer's feedback in this revision. This is retry {retry_count + 1} of {max_retries}."
        )
        # Remove any prior retry feedback blocks to avoid accumulation
        original_prompt = upstream_phase.get("system_prompt", "")
        # Strip previous retry blocks if present
        import re as _re
        original_prompt = _re.sub(
            r'\n\n# Retry attempt \d+/\d+ — Reviewer feedback from previous attempt:.*?'
            r'This is retry \d+ of \d+\.',
            '', original_prompt, flags=re.DOTALL
        )
        phases[upstream_phase_index]["system_prompt"] = original_prompt + feedback_block

        # Update pipeline: rewind to upstream phase
        p.phases = phases
        p.current_phase_index = upstream_phase_index
        p.status = "running"
        await db.commit()

    new_retry = retry_count + 1
    feedback_summary = feedback_text[:200] + ("..." if len(feedback_text) > 200 else "")
    _push(
        pipeline_id, "phase_retry",
        phase_index=upstream_phase_index,
        phase_name=phase_name,
        retry_count=new_retry,
        max_retries=max_retries,
        feedback_summary=feedback_summary,
        message=f"Retrying phase '{phase_name}' with reviewer feedback (attempt {new_retry}/{max_retries}).",
    )

    logger.info(
        f"Pipeline {pipeline_id}: retrying phase {upstream_phase_index} ({phase_name}) "
        f"with feedback — attempt {new_retry}/{max_retries}"
    )

    # Run the upstream phase again
    asyncio.create_task(_run_phase(pipeline_id, upstream_phase_index, retry_count=new_retry, max_retries=max_retries))
    return True


# ─── Routes ───────────────────────────────────────────────────────────────────
@router.get("/methods/{method_id}/phases", dependencies=[Depends(verify_api_key)])
async def preview_phases(method_id: str):
    """Return the ordered list of phases for a method (without starting a pipeline).

    For each phase, reports the full Phase → Agent → Persona → Model chain
    so the frontend can show the user what will actually run.
    """
    phase_rows = await _build_phase_preview_rows(method_id)

    return {
        "method_id": method_id,
        "phases": phase_rows,
    }


@router.post("/methods/preview", dependencies=[Depends(verify_api_key)])
async def preview_phases_with_stack(body: PipelinePhasePreviewRequest):
    """Stack-aware runtime preview used by the launcher before pipeline creation."""
    effective_method_id, requested_stack, stacked_prompt, layered_methods = _resolve_pipeline_method_selection(
        body.method_id,
        body.stack_override,
    )
    phase_rows = await _build_phase_preview_rows(effective_method_id)

    return {
        "method_id": body.method_id,
        "effective_method_id": effective_method_id,
        "stack_override": requested_stack,
        "layered_methods": layered_methods,
        "stacked_prompt_applied": bool(stacked_prompt),
        "phases": phase_rows,
    }


@router.post("", dependencies=[Depends(verify_api_key)])
async def create_pipeline(body: PipelineCreate, request: Request, db: AsyncSession = Depends(get_db)):
    """Create and start a multi-agent pipeline attached to a workbench session."""
    from app.models.pipeline import Pipeline
    from app.models.workbench import WorkbenchSession
    from app.services.phase_templates import get_phases_for_method, list_supported_methods, validate_phase_dag, get_method_phases_with_custom

    # Resolve method — if "stack" or "active", use the currently-active method
    # stack from the Methods page. The primary (first) method in the stack
    # determines the phase structure; other stacked methods' prompts get
    # injected into every phase.
    effective_method_id, requested_stack, stacked_prompt, layered_methods = _resolve_pipeline_method_selection(
        body.method_id,
        body.stack_override,
    )
    if layered_methods:
        logger.info(f"Pipeline method={effective_method_id}, layering stack: {layered_methods}")

    # Validate method — check custom methods (DB) first, then built-in
    try:
        template_phases = await get_method_phases_with_custom(effective_method_id, db)
    except KeyError:
        raise HTTPException(
            status_code=400,
            detail=f"Method '{effective_method_id}' does not support pipelines. "
                   f"Supported: {list_supported_methods()}"
        )

    # Validate phase dependency DAG
    dag_errors = validate_phase_dag(template_phases)
    if dag_errors:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid phase dependency graph: {'; '.join(dag_errors)}"
        )

    # Validate session
    session = (await db.execute(
        select(WorkbenchSession).where(WorkbenchSession.id == body.session_id)
    )).scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail=f"Workbench session {body.session_id} not found")

    # Warn if no project is attached — code phases will produce artifacts but
    # files won't land on disk. Not a hard failure because pure analysis
    # pipelines can still be useful without a project.
    if not session.project_path:
        logger.warning(
            f"Pipeline created for session {body.session_id} WITHOUT a project_path. "
            f"Code phases will produce artifacts but files won't be written to disk. "
            f"Attach the session to a project for the full workflow."
        )

    # Apply only explicit per-phase model overrides from the user.
    # If no override is provided, keep phase model unset so runtime resolution
    # can follow phase binding -> agent -> persona -> template default.
    overrides = body.model_overrides or {}
    for phase in template_phases:
        if phase["name"] in overrides:
            phase["model"] = overrides[phase["name"]]
        else:
            phase.pop("model", None)
        # Inject stacked method prompts into each phase so secondary methods
        # (e.g., GTrack "commit after every change") apply to all phases.
        if stacked_prompt:
            existing = phase.get("system_prompt", "")
            phase["system_prompt"] = f"{existing}\n\n---\n\n# Additional method instructions (from stack)\n{stacked_prompt}"

    template_phases = _apply_interaction_mode_to_phases(
        template_phases,
        interaction_mode=body.interaction_mode,
        delegate_qa_to_agent=body.delegate_qa_to_agent,
    )

    effective_auto_approve = body.auto_approve
    if (body.interaction_mode or "autonomous").strip().lower() == "interactive":
        # Interactive mode must pause at each phase handoff for user engagement.
        effective_auto_approve = False

    # Determine creator user ID
    user = getattr(request.state, "user", None)
    creator_id = user.get("id", "owner") if user else "owner"

    # Validate approval_policy
    policy = (body.approval_policy or "any").lower()
    if policy not in ("any", "majority", "all"):
        raise HTTPException(status_code=400, detail=f"Invalid approval_policy '{policy}'. Must be 'any', 'majority', or 'all'.")

    # Create Pipeline row
    pipeline_id = str(uuid.uuid4())
    pipeline = Pipeline(
        id=pipeline_id,
        session_id=body.session_id,
        method_id=effective_method_id,
        phases=template_phases,
        current_phase_index=0,
        status="pending",
        auto_approve=effective_auto_approve,
        approvers=body.approvers or [],
        approval_policy=policy,
        created_by=creator_id,
        initial_task=body.task,
    )
    db.add(pipeline)

    # Link session -> pipeline
    session.pipeline_id = pipeline_id
    await db.commit()

    # Set up SSE queue
    _queues[pipeline_id] = asyncio.Queue(maxsize=1000)
    _event_logs[pipeline_id] = []

    _push(pipeline_id, "pipeline_created",
          method_id=body.method_id,
          phases=[{"name": p["name"], "role": p["role"], "model": p.get("model")} for p in template_phases],
          auto_approve=effective_auto_approve,
          interaction_mode=(body.interaction_mode or "autonomous").strip().lower(),
          delegate_qa_to_agent=bool(body.delegate_qa_to_agent))

    # Kick off all root phases (phases with no dependencies)
    await _advance_to_next(pipeline_id)

    return pipeline.to_dict()


@router.get("", dependencies=[Depends(verify_api_key)])
async def list_pipelines(db: AsyncSession = Depends(get_db)):
    from app.models.pipeline import Pipeline
    result = await db.execute(
        select(Pipeline).order_by(desc(Pipeline.created_at)).limit(100)
    )
    pipelines = result.scalars().all()
    return {"data": [p.to_dict() for p in pipelines], "total": len(pipelines)}


@router.get("/{pipeline_id}", dependencies=[Depends(verify_api_key)])
async def get_pipeline(pipeline_id: str, db: AsyncSession = Depends(get_db)):
    from app.models.pipeline import Pipeline, PhaseRun
    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    runs = (await db.execute(
        select(PhaseRun).where(PhaseRun.pipeline_id == pipeline_id).order_by(PhaseRun.phase_index, PhaseRun.created_at)
    )).scalars().all()

    return {
        **p.to_dict(),
        "phase_runs": [r.to_dict() for r in runs],
    }


@router.post("/{pipeline_id}/message", dependencies=[Depends(verify_api_key)])
async def send_pipeline_message(
    pipeline_id: str,
    body: PipelineMessage,
    db: AsyncSession = Depends(get_db),
):
    """Add user guidance that will be visible in the live activity feed and next phase context."""
    from app.models.pipeline import Pipeline
    from app.models.workbench import WorkbenchSession

    message = (body.message or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    pipeline = (await db.execute(
        select(Pipeline).where(Pipeline.id == pipeline_id)
    )).scalar_one_or_none()
    if not pipeline:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if pipeline.status in ("completed", "failed", "cancelled"):
        raise HTTPException(status_code=409, detail=f"Pipeline is already {pipeline.status}")

    session = (await db.execute(
        select(WorkbenchSession).where(WorkbenchSession.id == pipeline.session_id)
    )).scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Workbench session not found")

    applies_to = (
        "current approval or next retry"
        if pipeline.status == "awaiting_approval"
        else "next phase"
    )

    messages = list(session.messages or [])
    messages.append({
        "role": "user",
        "content": message,
        "kind": "pipeline_note",
        "created_at": datetime.utcnow().isoformat(),
        "applies_to": applies_to,
    })
    session.messages = messages
    await db.commit()

    _push(
        pipeline_id,
        "user_message",
        message=message,
        applies_to=applies_to,
        pipeline_status=pipeline.status,
    )
    return {"ok": True, "message": message, "applies_to": applies_to}


@router.post("/{pipeline_id}/approve", dependencies=[Depends(verify_api_key)])
async def approve_phase(
    pipeline_id: str,
    body: PipelineApprove,
    request: Request,
    phase_index: Optional[int] = None,
    phase_run_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    """Approve a phase — supports multi-approver collaborative approval.

    Records the individual approval gesture. If the approval threshold
    (based on pipeline's approval_policy) is met, advances the pipeline.
    With parallel execution multiple phases may be awaiting approval
    simultaneously.  Use ``phase_index`` or ``phase_run_id`` query params to
    target a specific phase; if omitted, falls back to the first
    awaiting_approval run found (backward-compat).
    """
    from app.models.pipeline import Pipeline, PhaseRun
    from app.services.ws_manager import manager

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    # Resolve approver user ID
    user = getattr(request.state, "user", None)
    approver_id = body.user_id or (user.get("id", "owner") if user else "owner")

    # Build query for the awaiting_approval PhaseRun
    q = (
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.status == "awaiting_approval")
    )
    if phase_run_id:
        q = q.where(PhaseRun.id == phase_run_id)
    elif phase_index is not None:
        q = q.where(PhaseRun.phase_index == phase_index)
    else:
        # Backward-compat: try current_phase_index first, then any
        q = q.order_by(PhaseRun.phase_index)
    q = q.order_by(desc(PhaseRun.created_at)).limit(1)

    pr = (await db.execute(q)).scalar_one_or_none()
    if not pr:
        raise HTTPException(status_code=404, detail="No phase awaiting approval")

    approved_index = pr.phase_index

    # Get approvers list and policy
    approvers = p.approvers or []
    policy = p.approval_policy or "any"

    # Record individual approval
    current_approvals = list(pr.approvals or [])

    # Prevent duplicate votes
    already_voted = any(a["user_id"] == approver_id for a in current_approvals)
    if already_voted:
        raise HTTPException(status_code=409, detail=f"User '{approver_id}' has already voted on this phase")

    current_approvals.append({
        "user_id": approver_id,
        "action": "approve",
        "feedback": body.feedback,
        "timestamp": datetime.utcnow().isoformat(),
    })
    pr.approvals = current_approvals

    # Push WebSocket event for real-time approval tracking
    _push(pipeline_id, "phase_approval_update",
          phase_index=approved_index,
          approver=approver_id,
          action="approve",
          approvals=current_approvals,
          feedback=body.feedback)

    try:
        await manager.broadcast_to_channel(f"pipeline:{pipeline_id}", {
            "type": "pipeline_approval_update",
            "payload": {
                "pipeline_id": pipeline_id,
                "phase_index": approved_index,
                "approver": approver_id,
                "action": "approve",
                "approvals": current_approvals,
            },
        })
    except Exception as ws_err:
        logger.debug(f"WebSocket broadcast for approval failed: {ws_err}")

    # Check threshold
    effective_approvers = approvers if approvers else [p.created_by or "owner"]
    threshold_result = _check_approval_threshold(current_approvals, effective_approvers, policy)

    if threshold_result == "approved":
        pr.status = "approved"
        if body.feedback:
            pr.user_feedback = body.feedback
        await db.commit()

        _push(pipeline_id, "phase_approved", phase_index=approved_index, feedback=body.feedback)
        await _advance_to_next(pipeline_id)

        return {"ok": True, "approved_phase": approved_index, "threshold_met": True,
                "approvals": current_approvals, "policy": policy}
    else:
        # Threshold not yet met — save and wait for more approvals
        await db.commit()
        return {"ok": True, "approved_phase": approved_index, "threshold_met": False,
                "approvals": current_approvals, "policy": policy,
                "pending_approvers": [uid for uid in effective_approvers
                                       if not any(a["user_id"] == uid for a in current_approvals)]}


@router.post("/{pipeline_id}/reject", dependencies=[Depends(verify_api_key)])
async def reject_phase(
    pipeline_id: str,
    body: PipelineReject,
    request: Request,
    phase_index: Optional[int] = None,
    phase_run_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    """Reject a phase with feedback — re-runs the same phase.

    Records the individual rejection in the approvals list. Any single
    rejection immediately rejects the phase.
    Supports ``phase_index`` or ``phase_run_id`` query params to target a
    specific phase when multiple are awaiting approval in parallel.
    """
    from app.models.pipeline import Pipeline, PhaseRun
    from app.services.ws_manager import manager

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    # Resolve rejector user ID
    user = getattr(request.state, "user", None)
    rejector_id = user.get("id", "owner") if user else "owner"

    q = (
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.status == "awaiting_approval")
    )
    if phase_run_id:
        q = q.where(PhaseRun.id == phase_run_id)
    elif phase_index is not None:
        q = q.where(PhaseRun.phase_index == phase_index)
    else:
        q = q.order_by(PhaseRun.phase_index)
    q = q.order_by(desc(PhaseRun.created_at)).limit(1)

    pr = (await db.execute(q)).scalar_one_or_none()
    if not pr:
        raise HTTPException(status_code=404, detail="No phase awaiting approval")

    rejected_index = pr.phase_index

    # Record individual rejection in approvals list
    current_approvals = list(pr.approvals or [])
    current_approvals.append({
        "user_id": rejector_id,
        "action": "reject",
        "feedback": body.feedback,
        "timestamp": datetime.utcnow().isoformat(),
    })
    pr.approvals = current_approvals
    pr.status = "rejected"
    pr.user_feedback = body.feedback
    await db.commit()

    _push(pipeline_id, "phase_rejected", phase_index=rejected_index, feedback=body.feedback,
          rejector=rejector_id, approvals=current_approvals)

    # Push WebSocket event
    try:
        await manager.broadcast_to_channel(f"pipeline:{pipeline_id}", {
            "type": "pipeline_approval_update",
            "payload": {
                "pipeline_id": pipeline_id,
                "phase_index": rejected_index,
                "approver": rejector_id,
                "action": "reject",
                "approvals": current_approvals,
            },
        })
    except Exception as ws_err:
        logger.debug(f"WebSocket broadcast for rejection failed: {ws_err}")

    # Bake the rejection feedback into the phase's system prompt for the re-run
    try:
        phases = list(p.phases or [])
        if 0 <= rejected_index < len(phases):
            original = phases[rejected_index].get("system_prompt", "")
            feedback_block = (f"\n\n# User rejected your previous attempt — feedback:\n"
                              f"{body.feedback}\n\nAddress this feedback in your next attempt.")
            if feedback_block not in original:
                phases[rejected_index]["system_prompt"] = original + feedback_block
            await _db_update_pipeline(pipeline_id, phases=phases, status="running")
    except Exception as e:
        logger.warning(f"Could not inject rejection feedback: {e}")

    # Re-run the same phase
    asyncio.create_task(_run_phase(pipeline_id, rejected_index))

    return {"ok": True, "re_running_phase": rejected_index}


@router.post("/{pipeline_id}/skip", dependencies=[Depends(verify_api_key)])
async def skip_phase(
    pipeline_id: str,
    body: PipelineSkip,
    phase_index: Optional[int] = None,
    phase_run_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    """Skip a phase and advance the pipeline.

    Supports ``phase_index`` or ``phase_run_id`` query params to target a
    specific phase when multiple are running in parallel.
    """
    from app.models.pipeline import Pipeline, PhaseRun
    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status not in ("awaiting_approval", "running"):
        raise HTTPException(status_code=409,
                            detail=f"Pipeline cannot skip from status={p.status}")

    # Find the phase run to skip
    q = (
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.status.in_(("running", "awaiting_approval")))
    )
    if phase_run_id:
        q = q.where(PhaseRun.id == phase_run_id)
    elif phase_index is not None:
        q = q.where(PhaseRun.phase_index == phase_index)
    else:
        q = q.order_by(PhaseRun.phase_index)
    q = q.order_by(desc(PhaseRun.created_at)).limit(1)

    pr = (await db.execute(q)).scalar_one_or_none()
    skipped_index = pr.phase_index if pr else (phase_index or p.current_phase_index)
    if pr:
        pr.status = "skipped"
        if body.reason:
            pr.user_feedback = body.reason
        await db.commit()

    _push(pipeline_id, "phase_skipped", phase_index=skipped_index, reason=body.reason)

    await _advance_to_next(pipeline_id)
    return {"ok": True, "skipped_phase": skipped_index}


@router.get("/{pipeline_id}/approvals", dependencies=[Depends(verify_api_key)])
async def get_approval_status(
    pipeline_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Show approval status for each phase — who approved, who is pending."""
    from app.models.pipeline import Pipeline, PhaseRun

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    runs = (await db.execute(
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .order_by(PhaseRun.phase_index, PhaseRun.created_at)
    )).scalars().all()

    approvers = p.approvers or []
    effective_approvers = approvers if approvers else [p.created_by or "owner"]
    policy = p.approval_policy or "any"

    phases_status = []
    for run in runs:
        run_approvals = run.approvals or []
        voted_users = {a["user_id"] for a in run_approvals}
        pending_users = [uid for uid in effective_approvers if uid not in voted_users]

        phases_status.append({
            "phase_index": run.phase_index,
            "phase_name": run.phase_name,
            "phase_run_id": run.id,
            "status": run.status,
            "approvals": run_approvals,
            "pending_approvers": pending_users if run.status == "awaiting_approval" else [],
            "threshold_met": run.status == "approved",
        })

    return {
        "pipeline_id": pipeline_id,
        "approval_policy": policy,
        "approvers": effective_approvers,
        "phases": phases_status,
    }


@router.post("/{pipeline_id}/retry", dependencies=[Depends(verify_api_key)])
async def retry_phase(
    pipeline_id: str,
    phase_index: Optional[int] = None,
    db: AsyncSession = Depends(get_db),
):
    """Retry failed phase(s).

    If ``phase_index`` is given, retry that specific phase.  Otherwise retry
    all failed phases whose dependencies are met (enables parallel retry).
    """
    from app.models.pipeline import Pipeline, PhaseRun
    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status not in ("failed", "cancelled"):
        raise HTTPException(status_code=409,
                            detail=f"Can only retry from failed/cancelled (status={p.status})")

    # Reset pipeline status
    await _db_update_pipeline(pipeline_id, status="running", completed_at=None)

    phases = list(p.phases or [])

    if phase_index is not None:
        # Retry a single specific phase
        phase_name = phases[phase_index].get("name") if 0 <= phase_index < len(phases) else f"Phase {phase_index + 1}"
        _push(
            pipeline_id,
            "pipeline_retry",
            phase_index=phase_index,
            phase_name=phase_name,
            retried_indices=[phase_index],
            message=f"Retrying {phase_name}.",
        )
        asyncio.create_task(_run_phase(pipeline_id, phase_index))
        return {"ok": True, "retrying_phase": phase_index}

    # Retry all failed phases whose deps are met — let _advance_to_next
    # figure it out after we clear the failed PhaseRun records so they
    # don't block re-evaluation.
    failed_runs = (await db.execute(
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.status == "failed")
    )).scalars().all()

    retried_indices: list[int] = []
    for fr in failed_runs:
        # Delete the failed run so _advance_to_next sees the phase as
        # pending and will re-launch it if deps are met.
        await db.delete(fr)
        retried_indices.append(fr.phase_index)
    await db.commit()

    unique_indices = sorted(set(retried_indices))
    phase_labels = [
        phases[idx].get("name") if 0 <= idx < len(phases) else f"Phase {idx + 1}"
        for idx in unique_indices
    ]
    if not unique_indices:
        retry_message = "No failed phases were waiting to be retried. Resuming the pipeline."
    elif len(unique_indices) == 1:
        retry_message = f"Retrying {phase_labels[0]}."
    elif len(unique_indices) <= 4:
        retry_message = f"Retrying phases: {', '.join(phase_labels)}."
    else:
        retry_message = f"Retrying {len(unique_indices)} failed phases."

    _push(
        pipeline_id,
        "pipeline_retry",
        phase_index=unique_indices[0] if unique_indices else 0,
        phase_name=phase_labels[0] if phase_labels else None,
        retried_indices=unique_indices,
        retried_phase_names=phase_labels,
        message=retry_message,
    )

    await _advance_to_next(pipeline_id)
    return {"ok": True, "retrying_phases": unique_indices, "retrying_phase_names": phase_labels}


@router.post("/{pipeline_id}/phase-model", dependencies=[Depends(verify_api_key)])
async def update_phase_model(
    pipeline_id: str,
    body: PipelinePhaseModelUpdate,
    db: AsyncSession = Depends(get_db),
):
    """Change the configured model for one phase and immediately rerun that phase.

    This is intended for recovery when the current assigned model is invalid,
    unavailable, or simply not the right choice for the phase.
    """
    from app.models.pipeline import Pipeline, PhaseRun
    from app.routes.workbench import _resolve_model

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status == "completed":
        raise HTTPException(status_code=409, detail="Completed pipelines cannot change phase models")

    phases = list(p.phases or [])
    if body.phase_index < 0 or body.phase_index >= len(phases):
        raise HTTPException(status_code=400, detail="Invalid phase index")

    model_ref = (body.model_id or "").strip()
    if not model_ref:
        raise HTTPException(status_code=400, detail="Model is required")

    model_orm, _provider_orm = await _resolve_model(model_ref)
    if not model_orm:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Model '{model_ref}' is not active, validated, and ready for runtime use. "
                "Choose another confirmed model from the dropdown."
            ),
        )

    latest_run = (await db.execute(
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.phase_index == body.phase_index)
        .order_by(desc(PhaseRun.created_at))
        .limit(1)
    )).scalar_one_or_none()

    if latest_run and latest_run.status == "running":
        raise HTTPException(
            status_code=409,
            detail=(
                "This phase is actively running right now. Wait for it to finish or fail before switching models."
            ),
        )

    phase = dict(phases[body.phase_index] or {})
    previous_model = phase.get("model") or (latest_run.model_id if latest_run else None) or phase.get("default_model")
    phase["model"] = model_orm.model_id
    phases[body.phase_index] = phase
    p.phases = phases
    p.status = "running"
    p.completed_at = None
    p.current_phase_index = body.phase_index

    supersede_note = f"Superseded after manual model change to {model_orm.model_id}."
    runs_to_supersede = (await db.execute(
        select(PhaseRun)
        .where(PhaseRun.pipeline_id == pipeline_id)
        .where(PhaseRun.phase_index == body.phase_index)
        .where(PhaseRun.status.in_(("failed", "awaiting_approval")))
    )).scalars().all()

    for run in runs_to_supersede:
        run.status = "rejected"
        if not run.user_feedback:
            run.user_feedback = supersede_note

    await db.commit()

    phase_name = phase.get("name") or (latest_run.phase_name if latest_run else f"Phase {body.phase_index + 1}")
    _push(
        pipeline_id,
        "phase_model_changed",
        phase_index=body.phase_index,
        phase_name=phase_name,
        previous_model=previous_model,
        model_id=model_orm.model_id,
        message=(
            f"Phase '{phase_name}' switched from '{previous_model or 'unassigned'}' to "
            f"'{model_orm.model_id}'. Re-running with the updated model."
        ),
    )

    asyncio.create_task(_run_phase(pipeline_id, body.phase_index))
    return {
        "ok": True,
        "phase_index": body.phase_index,
        "phase_name": phase_name,
        "previous_model": previous_model,
        "model_id": model_orm.model_id,
        "re_running_phase": body.phase_index,
    }


@router.post("/{pipeline_id}/cancel", dependencies=[Depends(verify_api_key)])
async def cancel_pipeline(pipeline_id: str, db: AsyncSession = Depends(get_db)):
    from app.models.pipeline import Pipeline
    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    await _db_update_pipeline(pipeline_id, status="cancelled", completed_at=datetime.utcnow())
    _push(pipeline_id, "pipeline_done", message="Pipeline cancelled.", status="cancelled")
    return {"ok": True}


@router.post("/{pipeline_id}/pause", dependencies=[Depends(verify_api_key)])
async def pause_pipeline(pipeline_id: str, db: AsyncSession = Depends(get_db)):
    from app.models.pipeline import Pipeline

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status in ("completed", "failed", "cancelled"):
        raise HTTPException(status_code=409, detail=f"Cannot pause a pipeline in status '{p.status}'")
    if p.status == "paused":
        return {"ok": True, "status": "paused"}

    await _db_update_pipeline(pipeline_id, status="paused")
    _push(pipeline_id, "pipeline_paused", message="Pipeline paused by user.", status="paused")
    return {"ok": True, "status": "paused"}


@router.post("/{pipeline_id}/resume", dependencies=[Depends(verify_api_key)])
async def resume_pipeline(pipeline_id: str, db: AsyncSession = Depends(get_db)):
    from app.models.pipeline import Pipeline

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status != "paused":
        raise HTTPException(status_code=409, detail=f"Pipeline is not paused (current status: '{p.status}')")

    await _db_update_pipeline(pipeline_id, status="running")
    _push(pipeline_id, "pipeline_resumed", message="Pipeline resumed.", status="running")
    await _advance_to_next(pipeline_id)
    return {"ok": True, "status": "running"}


@router.delete("/{pipeline_id}", dependencies=[Depends(verify_api_key)])
async def delete_pipeline(pipeline_id: str, db: AsyncSession = Depends(get_db)):
    """Delete a single pipeline and all its phase runs. Session is NOT deleted."""
    from app.models.pipeline import Pipeline, PhaseRun
    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    # Delete phase runs
    runs = (await db.execute(select(PhaseRun).where(PhaseRun.pipeline_id == pipeline_id))).scalars().all()
    for r in runs:
        await db.delete(r)
    await db.delete(p)
    await db.commit()
    _queues.pop(pipeline_id, None)
    _event_logs.pop(pipeline_id, None)
    return {"ok": True, "deleted_phase_runs": len(runs)}


# ─── Save pipeline as method template ─────────────────────────────────────────

class SaveAsTemplateRequest(BaseModel):
    name: str
    description: Optional[str] = None
    include_system_prompts: bool = False


@router.post("/{pipeline_id}/save-as-template", dependencies=[Depends(verify_api_key)])
async def save_pipeline_as_template(
    pipeline_id: str,
    body: SaveAsTemplateRequest,
    db: AsyncSession = Depends(get_db),
):
    """Convert a completed pipeline's phase config into a reusable custom method template."""
    from app.models.pipeline import Pipeline
    from app.models.custom_method import CustomMethod

    p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    if p.status != "completed":
        raise HTTPException(status_code=400, detail="Only completed pipelines can be saved as templates")

    # Check for duplicate name
    existing = await db.execute(select(CustomMethod).where(CustomMethod.name == body.name))
    if existing.scalar_one_or_none():
        raise HTTPException(status_code=409, detail=f"A method named '{body.name}' already exists")

    # Convert pipeline phases to method template phases
    template_phases = []
    for phase in (p.phases or []):
        template_phase: Dict[str, Any] = {
            "name": phase.get("name", ""),
            "role": phase.get("role", ""),
            "artifact_type": phase.get("artifact_type", "md"),
        }
        # Preserve model overrides as defaults
        model = phase.get("model") or phase.get("default_model")
        if model:
            template_phase["default_model"] = model
        # Optionally include system prompts
        if body.include_system_prompts and phase.get("system_prompt"):
            template_phase["system_prompt"] = phase["system_prompt"]
        # Preserve depends_on if present
        if phase.get("depends_on"):
            template_phase["depends_on"] = phase["depends_on"]
        # Preserve conditions if present
        if phase.get("conditions"):
            template_phase["conditions"] = phase["conditions"]

        template_phases.append(template_phase)

    new_method = CustomMethod(
        name=body.name,
        description=body.description or f"Template created from pipeline run",
        phases=template_phases,
        is_active=True,
    )
    db.add(new_method)
    await db.commit()
    await db.refresh(new_method)

    return {"ok": True, "method": new_method.to_dict(), "source_pipeline_id": pipeline_id}


def _string_list(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


async def _load_discovery_handoff(pipeline_id: str) -> tuple[Any, dict]:
    from app.models.pipeline import Pipeline, PhaseRun

    async with AsyncSessionLocal() as db:
        pipeline = (await db.execute(
            select(Pipeline).where(Pipeline.id == pipeline_id)
        )).scalar_one_or_none()
        if not pipeline:
            raise HTTPException(status_code=404, detail="Pipeline not found")
        if pipeline.method_id != "discovery":
            raise HTTPException(status_code=400, detail="Discovery export is only available for discovery pipelines")

        run = (await db.execute(
            select(PhaseRun)
            .where(PhaseRun.pipeline_id == pipeline_id)
            .where(PhaseRun.phase_name == "HandoffPlanner")
            .order_by(desc(PhaseRun.created_at))
        )).scalars().first()
        if not run or not isinstance((run.output_artifact or {}).get("data"), dict):
            raise HTTPException(status_code=400, detail="Discovery handoff artifact is not available yet")

        return pipeline, run.output_artifact["data"]


def _build_discovery_export_lines(pipeline: Any, handoff: dict) -> List[str]:
    use_case = handoff.get("use_case") if isinstance(handoff.get("use_case"), dict) else {}
    requirements = handoff.get("requirements_snapshot") if isinstance(handoff.get("requirements_snapshot"), dict) else {}

    lines: List[str] = []
    lines.append(handoff.get("handoff_title") or "Discovery Handoff")
    lines.append("")
    lines.append(f"Generated: {datetime.utcnow().isoformat()} UTC")
    lines.append(f"Pipeline ID: {pipeline.id}")
    lines.append(f"Session ID: {pipeline.session_id}")
    lines.append("")

    summary = str(handoff.get("summary") or pipeline.initial_task or "").strip()
    if summary:
        lines.append("Summary")
        lines.append(summary)
        lines.append("")

    if use_case:
        lines.append("Use Case")
        actor = str(use_case.get("actor") or "").strip()
        need = str(use_case.get("need") or "").strip()
        value = str(use_case.get("value") or "").strip()
        if actor:
            lines.append(f"Actor: {actor}")
        if need:
            lines.append(f"Need: {need}")
        if value:
            lines.append(f"Value: {value}")
        lines.append("")

    functional = _string_list(requirements.get("functional"))
    non_functional = _string_list(requirements.get("non_functional"))
    acceptance = _string_list(requirements.get("acceptance_criteria"))
    open_decisions = _string_list(handoff.get("open_decisions"))
    recommended_stack = _string_list(handoff.get("recommended_stack"))
    next_phase_brief = str(handoff.get("next_phase_brief") or "").strip()

    if functional:
        lines.append("Functional Requirements")
        lines.extend([f"- {item}" for item in functional])
        lines.append("")

    if non_functional:
        lines.append("Non-Functional Requirements")
        lines.extend([f"- {item}" for item in non_functional])
        lines.append("")

    if acceptance:
        lines.append("Acceptance Criteria")
        lines.extend([f"- {item}" for item in acceptance])
        lines.append("")

    if open_decisions:
        lines.append("Open Decisions")
        lines.extend([f"- {item}" for item in open_decisions])
        lines.append("")

    lines.append("Recommended Next Method")
    lines.append(str(handoff.get("recommended_next_method") or "bmad"))
    if recommended_stack:
        lines.append(f"Recommended Stack: {' + '.join(recommended_stack)}")
    lines.append("")

    if next_phase_brief:
        lines.append("Next Phase Brief")
        lines.append(next_phase_brief)
        lines.append("")

    return lines


def _build_discovery_docx_bytes(pipeline: Any, handoff: dict) -> bytes:
    from docx import Document

    lines = _build_discovery_export_lines(pipeline, handoff)
    document = Document()
    document.add_heading(lines[0], level=0)

    section_titles = {
        "Summary",
        "Use Case",
        "Functional Requirements",
        "Non-Functional Requirements",
        "Acceptance Criteria",
        "Open Decisions",
        "Recommended Next Method",
        "Next Phase Brief",
    }

    for line in lines[1:]:
        if not line:
            document.add_paragraph("")
        elif not line.startswith("- ") and line in section_titles:
            document.add_heading(line, level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_discovery_pdf_bytes(pipeline: Any, handoff: dict) -> bytes:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    lines = _build_discovery_export_lines(pipeline, handoff)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    body_style = styles["BodyText"]
    bullet_style = styles["BodyText"]
    heading_style.textColor = HexColor("#0f4c5c")
    body_style.leading = 15
    bullet_style.leading = 15

    section_titles = {
        "Summary",
        "Use Case",
        "Functional Requirements",
        "Non-Functional Requirements",
        "Acceptance Criteria",
        "Open Decisions",
        "Recommended Next Method",
        "Next Phase Brief",
    }

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    story = [Paragraph(escape(lines[0]), title_style), Spacer(1, 0.2 * inch)]

    for line in lines[1:]:
        if not line:
            story.append(Spacer(1, 0.12 * inch))
        elif not line.startswith("- ") and line in section_titles:
            story.append(Paragraph(escape(line), heading_style))
            story.append(Spacer(1, 0.06 * inch))
        elif line.startswith("- "):
            story.append(Paragraph(f"&#8226; {escape(line[2:])}", bullet_style))
        else:
            story.append(Paragraph(escape(line).replace("\n", "<br/>"), body_style))

    doc.build(story)
    return buffer.getvalue()


@router.get("/{pipeline_id}/discovery-export", dependencies=[Depends(verify_api_key)])
async def export_discovery_handoff(pipeline_id: str, format: str = "docx"):
    pipeline, handoff = await _load_discovery_handoff(pipeline_id)
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "-", str(handoff.get("handoff_title") or "discovery-handoff").strip()).strip("-") or "discovery-handoff"
    requested = (format or "docx").strip().lower()

    if requested == "docx":
        content = _build_discovery_docx_bytes(pipeline, handoff)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        suffix = "docx"
    elif requested == "pdf":
        content = _build_discovery_pdf_bytes(pipeline, handoff)
        media_type = "application/pdf"
        suffix = "pdf"
    else:
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")

    filename = f"{safe_name}.{suffix}"
    return StreamingResponse(
        BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{pipeline_id}/stream")
async def stream_pipeline(pipeline_id: str, request: Request):
    """SSE stream — no auth (EventSource limitation)."""
    from app.models.pipeline import Pipeline, PhaseRun
    async with AsyncSessionLocal() as db:
        p = (await db.execute(select(Pipeline).where(Pipeline.id == pipeline_id))).scalar_one_or_none()
        if not p:
            raise HTTPException(status_code=404, detail="Pipeline not found")

        runs = (await db.execute(
            select(PhaseRun).where(PhaseRun.pipeline_id == pipeline_id).order_by(PhaseRun.phase_index, PhaseRun.created_at)
        )).scalars().all()
        init_payload = {**p.to_dict(), "phase_runs": [r.to_dict() for r in runs]}

    active_statuses = {"pending", "running", "paused", "awaiting_approval"}

    async def event_generator() -> AsyncGenerator[str, None]:
        yield f"data: {json.dumps({'type':'init','payload':init_payload})}\n\n"

        queue = _queues.get(pipeline_id)
        if not queue:
            # Pipeline exists but no live queue — replay stored events
            for evt in _event_logs.get(pipeline_id, []):
                yield f"data: {json.dumps(evt)}\n\n"
                await asyncio.sleep(0.02)
            # Terminal pipelines should emit one final done event and end.
            if p.status in ("completed", "failed", "cancelled"):
                yield f"data: {json.dumps({'type':'pipeline_done','payload':{'status':p.status}})}\n\n"
                return

            # Non-terminal pipelines must keep the stream alive. Otherwise the
            # browser auto-reconnects and the UI flickers connecting/reconnecting.
            if p.status in active_statuses:
                _queues[pipeline_id] = asyncio.Queue(maxsize=1000)
                queue = _queues[pipeline_id]
            else:
                # Unknown status: avoid infinite reconnect churn.
                return

        while True:
            if await request.is_disconnected():
                break
            try:
                evt = await asyncio.wait_for(queue.get(), timeout=30.0)
                yield f"data: {json.dumps(evt)}\n\n"
                if evt.get("type") == "pipeline_done":
                    break
            except asyncio.TimeoutError:
                yield f"data: {json.dumps({'type':'ping','payload':{}})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Access-Control-Allow-Origin": "*",
        }
    )
